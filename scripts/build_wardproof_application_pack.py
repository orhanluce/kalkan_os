from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path('output/wardproof_basvuru_paketi')
OUT.mkdir(parents=True, exist_ok=True)

NAVY = RGBColor(18, 42, 66)
BLUE = RGBColor(36, 91, 133)
GRAY = RGBColor(90, 98, 108)
LIGHT = 'EEF3F7'

def font(run, size=11, bold=False, color=None, italic=False):
    run.font.name = 'Calibri'
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), 'Calibri')
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), 'Calibri')
    run.font.size = Pt(size); run.bold = bold; run.italic = italic
    if color: run.font.color.rgb = color

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc.get_or_add_tcPr(); tcMar = tc.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar = OxmlElement('w:tcMar'); tc.append(tcMar)
    for k,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=OxmlElement('w:'+k); node.set(qn('w:w'),str(v)); node.set(qn('w:type'),'dxa'); tcMar.append(node)

def setup(doc):
    sec=doc.sections[0]; sec.page_width=Inches(8.5); sec.page_height=Inches(11)
    sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)
    sec.header_distance=sec.footer_distance=Inches(.49)
    styles=doc.styles
    n=styles['Normal']; n.font.name='Calibri'; n.font.size=Pt(11); n.font.color.rgb=RGBColor(32,38,44)
    n.paragraph_format.space_after=Pt(8); n.paragraph_format.line_spacing=1.22
    for name,size,color,before,after in [('Title',25,NAVY,0,8),('Heading 1',16,BLUE,16,8),('Heading 2',13,BLUE,12,6),('Heading 3',11.5,NAVY,9,4)]:
        s=styles[name]; s.font.name='Calibri'; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=color
        s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
    header=sec.header.paragraphs[0]; header.text='WARDPROOF.COM  |  BAŞVURU DOSYASI'; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    for r in header.runs: font(r,8.5,True,GRAY)
    footer=sec.footer.paragraphs[0]; footer.text='Kurucu çalışma nüshası • 21 Temmuz 2026'; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in footer.runs: font(r,8.5,False,GRAY)

def add_title(doc,title,subtitle=None):
    p=doc.add_paragraph(style='Title'); p.alignment=WD_ALIGN_PARAGRAPH.LEFT; font(p.add_run(title),25,True,NAVY)
    if subtitle:
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(18); font(p.add_run(subtitle),12,False,GRAY)

def add_p(doc,text,bold_lead=None):
    p=doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        font(p.add_run(bold_lead),11,True,NAVY); font(p.add_run(text[len(bold_lead):]),11)
    else: font(p.add_run(text),11)
    return p

def bullet(doc,text):
    p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.left_indent=Inches(.5); p.paragraph_format.first_line_indent=Inches(-.25)
    p.paragraph_format.space_after=Pt(4); font(p.add_run(text),11); return p

def callout(doc,label,text):
    t=doc.add_table(rows=1,cols=1); t.autofit=False; t.alignment=WD_TABLE_ALIGNMENT.LEFT; t.columns[0].width=Inches(6.5)
    c=t.cell(0,0); shade(c,LIGHT); margins(c,150,180,150,180); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p=c.paragraphs[0]; font(p.add_run(label+'\n'),10,True,BLUE); font(p.add_run(text),11)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)

def page(doc): doc.add_page_break()

doc=Document(); setup(doc)
add_title(doc,'WARDPROOF.COM','Kuluçka, BiGG ve banka pilotu başvuruları için ana dosya')
callout(doc,'BU DOSYA NASIL KULLANILIR','Köşeli parantez içindeki kişisel alanları siz doldurun. Müşteri, gelir, ekip veya yatırım bilgisi yoksa boşluğu gerçek durumla tamamlayın; tahminî rakamı gerçekleşmiş gibi yazmayın. Programlara özel cevapları ilgili forma göre kısaltabilirsiniz.')
add_p(doc,'Proje adı: Wardproof.com', 'Proje adı:')
add_p(doc,'Teknik kod adı: KALKAN_OS', 'Teknik kod adı:')
add_p(doc,'Aşama: wardproof.com üzerinden erişilebilen çalışan ürün / ticari doğrulama öncesi', 'Aşama:')
add_p(doc,'Kurucu: [AD SOYAD]', 'Kurucu:')
add_p(doc,'İletişim: [E-POSTA] • [TELEFON] • wardproof.com', 'İletişim:')

doc.add_heading('Başvuru öncesi doldurulacak alanlar',1)
for x in ['Kurucunun kısa özgeçmişi ve LinkedIn adresi','Şirketleşme durumu ve varsa mevcut şirket ortaklıkları','Ekipte aktif çalışan kişiler ve haftalık ayırdıkları süre','Bugüne kadar yapılan müşteri/uzman görüşmesi sayısı','Varsa pilot, LOI, referans veya gelir bilgisi','Başvuruda istenen destek tutarı ve kullanım planı','Demo kullanıcı bilgisi veya iki dakikalık demo videosu']:
    bullet(doc,'[ ] '+x)

page(doc); add_title(doc,'1. Kısa anlatım bankası')
doc.add_heading('Tek cümlelik anlatım',1)
add_p(doc,'Wardproof, finans kuruluşlarının mevzuat yükümlülüklerini kontrol, test, kanıt, bulgu ve bağımsız kapanış adımlarına bağlayan bir sürekli uyum yazılımıdır.')
doc.add_heading('50 kelimelik anlatım',1)
add_p(doc,'Finans kurumlarında uyum çalışmaları çoğu zaman e-posta, Excel, dosya klasörleri ve farklı ekipler arasında yürütülüyor. Wardproof bu dağınık süreci tek yerde topluyor; hangi yükümlülüğün hangi kontrolle karşılandığını, kontrolün gerçekten test edilip edilmediğini ve kararın hangi kanıta dayandığını izlenebilir hale getiriyor.')
doc.add_heading('150 kelimelik anlatım',1)
add_p(doc,'Wardproof’u geliştirirken çıkış noktam şuydu: Bir kurumun “uyumluyuz” demesi ile bunu denetimde gösterebilmesi aynı şey değil. Mevzuat maddesi bir yerde, kontrol listesi başka yerde, kanıtlar klasörlerde, bulgular ise e-posta ve tablolarda kalıyor. Bu parçalı yapı hem zaman kaybettiriyor hem de kararın hangi dayanağa göre verildiğini belirsizleştiriyor. Wardproof, resmî kaynaktan başlayan yükümlülüğü uygulanabilirlik kararına, kontrole, teste, kanıta, bulguya ve yeniden teste kadar tek zincirde izliyor. Ürün bugün wardproof.com üzerinden erişilebilen çalışan bir ürün seviyesinde. Çok kiracılı yapı, yetki ayrımı, denetim kaydı, kanıt bütünlüğü, görevler ayrılığı, üçüncü taraf riski ve paylaşılabilir Proof Room gibi temel akışlar çalışıyor. Şimdi ihtiyacım yeni özellik eklemekten çok, ürünü dar bir pilot kapsamıyla gerçek bir finans kuruluşunda doğrulamak ve şirketleşme sürecini sağlıklı kurmak.')

doc.add_heading('Problem',1)
add_p(doc,'Finans kuruluşlarında uyum operasyonu hâlâ büyük ölçüde insan takibiyle ilerliyor. Kontrolün var olması, çalıştığının test edilmesi ve denetçiye savunulabilir kanıt sunulması arasında kopukluk oluşuyor. Sonuçta denetim hazırlığı uzuyor, aynı bilgi tekrar tekrar toplanıyor ve kritik kararların geçmişi kolayca kaybolabiliyor.')
doc.add_heading('Çözüm',1)
add_p(doc,'Wardproof her uyum iddiasını kaynağına ve yürütme kanıtına bağlayan bir çalışma alanı sunuyor. Sistem “ölçemedik” ile “başarısız oldu” durumlarını birbirine karıştırmıyor; kritik bulguların başarılı yeniden test ve bağımsız onay olmadan kapatılmasına izin vermiyor. Bu yaklaşım kuruma yalnızca görev listesi değil, denetlenebilir bir karar izi sağlıyor.')
doc.add_heading('Neden şimdi?',1)
add_p(doc,'Finans kurumlarının siber güvenlik, operasyonel dayanıklılık, üçüncü taraf riski ve yönetim sorumluluğuna ilişkin yükümlülükleri artıyor. Aynı dönemde kurumların kullandığı sistem ve tedarikçi sayısı da büyüyor. Bu nedenle dönemsel kontrol listeleri yerine sürekli, kanıta dayalı ve açıklanabilir bir uyum işletimine ihtiyaç oluşuyor.')

page(doc); add_title(doc,'2. Ürün ve teknoloji')
doc.add_heading('Bugün çalışan başlıca parçalar',1)
for x in ['Çok kiracılı kurum yapısı ve satır seviyesinde veri izolasyonu','Kontrol kütüphanesi, kanıt yükleme ve denetim izi','Kontrol test motoru; başarısız, bilinmiyor, eski ve istisna durumlarının ayrılması','Bulgu, yeniden test ve bağımsız kapanış akışı','Görevler ayrılığı (SoD), istisna, süre dolumu ve maker-checker','Üçüncü ve dördüncü taraf güvence akışları','DORA bilgi sicili, kritik hizmet ve bağımlılık görünümü','Paylaşılabilir Proof Room ve bağımsız doğrulama paketleri','Türkçe kullanıcı arayüzü ve erişilebilirlik kontrolleri']:
    bullet(doc,x)
doc.add_heading('Teknik yapı',1)
add_p(doc,'Uygulama Next.js ve TypeScript ile geliştirildi; veri katmanında Supabase/PostgreSQL kullanılıyor. Mimari belirli bir kapalı platforma bağımlı kalmayacak şekilde saf PostgreSQL ilkelerine yakın tutuluyor. Yetkilendirme yalnızca arayüzde değil, veritabanı politikaları ve değiştirilemezlik kurallarıyla da uygulanıyor.')
doc.add_heading('Doğrulama durumu',1)
add_p(doc,'21 Temmuz 2026 tarihli proje kayıtlarında yaklaşık 1.465 birim/RLS testi ve 76 gerçek tarayıcı akışı raporlanıyor. Bu sayılar ürünün ticari başarısını değil, teknik kapsam ve doğrulama disiplinini gösteriyor. Başvuru sırasında güncel temiz sürümde yeniden üretilerek kanıt olarak sunulmalı.')
doc.add_heading('Bilinçli sınırlar',1)
for x in ['Ürün hukuki görüş veya “uyum garantisi” vermez; karar insan onayıyla tamamlanır.','Henüz doğrulanmış ücretli müşteri, yenileme veya satış döngüsü verisi yoktur.','İlk üretim bağlantıları pilot müşterinin kullandığı sistemlere göre seçilecektir.','Kurumsal satın alma için bağımsız pentest, staging/restore provası ve anahtar yönetimi gibi kapılar tamamlanmalıdır.']:
    bullet(doc,x)

page(doc); add_title(doc,'3. İş modeli ve pazara giriş')
doc.add_heading('İlk müşteri profili',1)
add_p(doc,'İlk odak; SPK, BDDK, TCMB veya ilgili siber güvenlik ve operasyonel dayanıklılık yükümlülükleri altında çalışan finans kuruluşlarıdır. Özellikle uyum, iç denetim, bilgi güvenliği, operasyonel risk ve üçüncü taraf yönetimi ekipleri ilk kullanıcı grubunu oluşturur.')
doc.add_heading('İlk satış kaması',1)
add_p(doc,'İlk müşteriye bütün platformu satmak yerine tek bir denetim veya risk alanıyla girilecek. Önerilen başlangıç paketleri: 20–30 kritik kontrol için kanıt zinciri, DORA/üçüncü taraf güvence paketi veya görevler ayrılığı ve kritik ödeme süreci kontrolü.')
doc.add_heading('Gelir modeli',1)
add_p(doc,'Planlanan model yıllık B2B SaaS aboneliği ve kapsamı önceden belirlenmiş kurulum/pilot hizmetidir. Fiyat henüz müşteri tarafından doğrulanmadı. İlk pilotlarda amaç yüksek fiyat iddiası değil; kurulum süresini, kullanıcı emeğini ve denetim hazırlığına etkisini ölçmektir.')
doc.add_heading('Başarı göstergeleri',1)
for x in ['İlk değer anına kadar geçen süre','Tamamlanan kontrol–test–kanıt zinciri sayısı','Eksik veya eski kanıtların kapanma süresi','Denetim paketi hazırlama süresindeki değişim','Pilotun yıllık sözleşmeye dönüşmesi','Müşteri başına kurulum ve destek saati']:
    bullet(doc,x)
doc.add_heading('Rekabet farkı',1)
add_p(doc,'Wardproof’u “daha ucuz bir GRC” olarak konumlandırmıyorum. Fark, yerel yükümlülüğü teknik kontrol ve yürütme kanıtına bağlamak; başarısızlık, ölçülememe, eskime ve kabul edilmiş istisnayı ayrı tutmak; bulgu kapanışını yeniden test ve bağımsız onayla doğrulamaktır.')

page(doc); add_title(doc,'4. Kurucu anlatımı')
doc.add_heading('Bu projeyi neden yapıyorum?',1)
add_p(doc,'Uyum ürünlerinde bana en zayıf gelen nokta, çok sayıda kontrol ve rapor üretilmesine rağmen “bu karar hangi kaynağa ve hangi gerçek kanıta dayanıyor?” sorusunun her zaman kolay cevaplanamamasıydı. Wardproof’u bu soruya teknik ve denetlenebilir bir cevap verebilmek için geliştirdim. Önce gösterişli özelliklerden çok veri izolasyonu, değiştirilemez kayıt, bağımsız onay ve test edilebilir iş kurallarına odaklandım.')
doc.add_heading('Bugüne kadar ne yaptım?',1)
add_p(doc,'Ürünü çalışan bir web uygulamasına dönüştürdüm; gerçek veritabanı, kimlik doğrulama, yetkilendirme, test ve canlı dağıtım akışlarını kurdum. Kontrol–kanıt–bulgu–yeniden test zincirinin yanında görevler ayrılığı, üçüncü taraf güvence, kritik hizmet, Proof Room ve denetim paketleri gibi kurumsal akışları da uçtan uca çalışır hale getirdim.')
doc.add_heading('Şimdi neden destek arıyorum?',1)
add_p(doc,'Teknik olarak ilerleyebiliyorum; fakat şirketleşme, hukuk/regülasyon doğrulaması, kurumsal satış ve pilot erişimi aynı anda tek kişinin taşıyabileceği alanlar değil. Aradığım destek yalnız para değil. Doğru müşteriyle dar bir pilot, deneyimli mentorlar, şirketleşme rehberliği ve güvenilir bir iş ağı benim için en az finansman kadar değerli.')
doc.add_heading('Ekip',1)
add_p(doc,'Şu anda [TEK KURUCUYUM / EKİP BİLGİSİ]. Ürün ve teknik geliştirme tarafını ben yürütüyorum. Pilot aşamasında regülasyon/hukuk doğrulaması ve kurumsal satış tarafında tamamlayıcı kişilerle çalışmayı planlıyorum. Bu rolleri varmış gibi göstermiyorum; başvuracağım programın doğru ekip kurulumunda da destek olmasını bekliyorum.')

page(doc); add_title(doc,'5. Altı haftalık pilot önerisi')
doc.add_heading('Pilotun amacı',1)
add_p(doc,'Bir finans kuruluşunda sınırlı sayıdaki kritik yükümlülük ve kontrol için mevcut hazırlık sürecini Wardproof üzerinden yürütmek; zaman, izlenebilirlik ve kanıt kalitesindeki değişimi ölçmek.')
doc.add_heading('Önerilen kapsam',1)
for x in ['Tek kurum ve 5–15 kullanıcı','20–30 kritik kontrol veya tek bir üçüncü taraf güvence paketi','Mevcut kanıtların sınırlı ve kontrollü aktarımı','En fazla iki manuel veri kaynağı; üretim sistemlerine yazma yok','Haftalık kısa çalışma toplantısı','Pilot sonunda yönetim özeti ve doğrulanabilir kanıt paketi']:
    bullet(doc,x)
doc.add_heading('Haftalara göre çalışma',1)
for h,t in [('1','Kapsam, roller, başarı ölçütleri ve mevcut süreç fotoğrafı'),('2','Kontrol/yükümlülük eşlemesi ve kullanıcı kurulumu'),('3','Kanıt toplama ve ilk test koşuları'),('4','Eksik kanıtlar, bulgular ve sorumlu atamaları'),('5','Yeniden test, bağımsız onay ve denetim paketi'),('6','Önce/sonra ölçümü, karar ve yıllık kullanım önerisi')]:
    add_p(doc,f'{h}. hafta: {t}',f'{h}. hafta:')
doc.add_heading('Başarı ölçütleri',1)
for x in ['Pilot kapsamındaki kontrollerin en az %80’inde kaynak–kontrol–test–kanıt zincirinin kurulması','Eksik/ölçülemeyen durumların başarısız sonuçlardan açıkça ayrılması','En az bir bulgunun yeniden test ve bağımsız onay akışından geçirilmesi','Denetim paketi hazırlama süresinin başlangıç yöntemiyle karşılaştırılması','Kurumun güvenlik, hukuk ve satın alma gereksinimlerinin kayıt altına alınması']:
    bullet(doc,x)
callout(doc,'PİLOT SINIRI','Wardproof üretim sistemlerinde işlem başlatmaz, saldırı simülasyonu yapmaz ve hukuki uygunluk kararı vermez. Pilot; okuma, manuel kanıt, iş akışı ve doğrulama paketiyle sınırlıdır.')

page(doc); add_title(doc,'6. Destek talebi ve kaynak kullanımı')
doc.add_heading('Talep ettiğim destek',1)
add_p(doc,'Öncelikli ihtiyacım şirket kuruluşunu ve ilk kurumsal pilotu finanse edecek başlangıç kaynağıdır. Bunun yanında finans kuruluşlarına erişim, regülasyon uzmanlığı, B2B satış mentorluğu, hukuk/muhasebe desteği ve güvenlik doğrulaması bekliyorum.')
doc.add_heading('Kaynak kullanım öncelikleri',1)
for x in ['Şirket kuruluşu, muhasebe ve temel hukuki sözleşmeler','Staging, yedekleme/restore ve güvenli anahtar yönetimi','Bağımsız pentest ve kurumsal güven paketi','Regülasyon içeriğinin uzman doğrulaması','Pilot kurulum ve müşteri destek kapasitesi','İlk iki üretim bağlantısı','Sınırlı kurucu geçim payı ve temel operasyon giderleri']:
    bullet(doc,x)
add_p(doc,'Kesin bütçe, başvurulan programın destek kalemlerine göre hazırlanacaktır. Uygun olmayan gideri sonradan programa uydurmak yerine her kalem baştan yazılı olarak teyit edilecektir.')

page(doc); add_title(doc,'7. Programa özel cevaplar')
doc.add_heading('Albaraka Garaj',1)
doc.add_heading('Programdan beklentim',2)
add_p(doc,'Wardproof için Albaraka Garaj’a yalnız yatırım bulmak amacıyla başvurmuyorum. En değerli fırsatın, ürünü gerçek bir finans kurumu ihtiyacına göre daraltmak ve Albaraka ekosisteminde ölçülebilir bir PoC yapmak olduğunu düşünüyorum. Özellikle üçüncü taraf güvence, kontrol kanıtı ve görevler ayrılığı alanlarında kısa bir pilot tasarlamak; hukuk, muhasebe ve kurumsal satış tarafında mentorluk almak istiyorum.')
doc.add_heading('Neden Albaraka?',2)
add_p(doc,'Ürünün hedef kullanıcısı finans kuruluşları olduğu için genel bir hızlandırıcıdan çok, bankanın gerçek ekiplerine ve karar süreçlerine erişebileceğim bir yapı benim için daha anlamlı. Wardproof’un güçlü teknik tarafını gerçek kullanıcı ihtiyacı, satın alma süreci ve ticari kapsamla dengelemek istiyorum.')

doc.add_heading('CYBERPARK BİGG MARKA / TÜBİTAK BiGG',1)
doc.add_heading('Ar-Ge ve yenilik yönü',2)
add_p(doc,'Projenin Ar-Ge yönü yalnız mevzuat metinlerini dijital ortama taşımak değildir. Farklı kaynaklardan gelen yükümlülük, kontrol, test ve kanıt kayıtlarının tenant güvenliği altında ilişkilendirilmesi; sonuç durumlarının deterministik kurallarla ayrıştırılması; değiştirilemez karar ve kanıt zincirinin doğrulanabilir paketlere dönüştürülmesi üzerinde çalışıyorum. Aynı girdinin aynı sonucu vermesi, karar gerekçesinin saklanması ve kritik kapanışların bağımsız kanıtla sınırlandırılması ürünün teknik çekirdeğini oluşturuyor.')
doc.add_heading('Ticarileşme planı',2)
add_p(doc,'İlk aşamada Türkiye’de düzenlemeye tabi finans kuruluşlarına dar pilot paketleriyle ulaşacağım. Pilot sonucu doğrulandıktan sonra yıllık abonelik modeline geçilecek. Yerel yükümlülük ve kanıt zinciriyle başlayan ürünün DORA, üçüncü taraf riski ve AI yönetişimi gibi bölgesel alanlara genişlemesi hedefleniyor.')
doc.add_heading('BiGG yatırımının etkisi',2)
add_p(doc,'Yatırım; şirket kuruluşunu, güvenlik kapılarını, uzman doğrulamasını ve ilk pilotları aynı takvim içinde yürütmemi sağlayacak. Kaynağı geniş özellik üretmek yerine, ticari doğrulama ve kurumsal güven için gerekli dar işlere ayırmayı planlıyorum.')

page(doc); add_title(doc,'7. Programa özel cevaplar — devam')
doc.add_heading('BTM Ön Kuluçka',1)
doc.add_heading('Neden BTM?',2)
add_p(doc,'Ürünü teknik olarak ileri bir noktaya getirdim; şimdi en büyük ihtiyacım iş modelini ve satış dilini gerçek şirketlerle doğrulamak. BTM’nin İTO ağı, komiteleri ve kurumsal bağlantıları sayesinde CFO, iç denetim, bilgi güvenliği ve uyum ekipleriyle görüşme yapabilmek benim için doğrudan değer yaratır. Program sonunda şirketleşmeye hazır, pilot kapsamı ve fiyatı netleşmiş bir girişim olmak istiyorum.')
doc.add_heading('Altı ay sonunda hedefim',2)
add_p(doc,'En az iki niyet mektubu, bir aktif pilot, doğrulanmış bir fiyat aralığı, standart pilot sözleşmesi ve kurucu dışında da anlatılabilen sade bir satış süreci oluşturmak.')

doc.add_heading('İTÜ Çekirdek',1)
doc.add_heading('Programdan beklentim',2)
add_p(doc,'İTÜ Çekirdek’te ürünün teknik kapsamını büyütmekten çok, yatırım yapılabilir ve satılabilir bir girişime dönüşmesini hedefliyorum. Mentorlarla pazar odağını daraltmak, finansal teknoloji alanında doğru pilotlara ulaşmak, BiGG sürecine hazırlanmak ve Big Bang ağı üzerinden yatırımcı geri bildirimi almak istiyorum.')
doc.add_heading('Neden bu ekip/proje?',2)
add_p(doc,'Wardproof fikir sunumundan ibaret değil; çalışan, canlıya alınmış ve yoğun test altyapısı bulunan bir ürün. Buna rağmen ticari doğrulamayı varmış gibi göstermiyorum. Teknik riski önemli ölçüde azaltılmış bu yapıyı, doğru müşteri ve iş modeli desteğiyle pazarda doğrulamak istiyorum.')

doc.add_heading('Fintech Factory',1)
doc.add_heading('Programdan beklentim',2)
add_p(doc,'Fintech Factory’den beklentim, ürünün finans kuruluşlarında gerçekten satın alınacak ilk kullanım alanını seçmek ve doğru karar vericilere ulaşmak. Çalışan demoyu bankacılık ve fintech profesyonellerinin geri bildirimiyle sınamak, pilot kapsamını sadeleştirmek ve ilk kurumsal referansı oluşturmak istiyorum.')
doc.add_heading('Girişimin aşaması',2)
add_p(doc,'Çalışan MVP mevcut; henüz ürün–pazar uyumu ve tekrarlanabilir gelir doğrulanmadı. Bu nedenle aşamayı “MVP / gelir öncesi ürün–pazar uyumu arıyor” olarak işaretleyeceğim.')

page(doc); add_title(doc,'7. Programa özel cevaplar — banka pilotları')
doc.add_heading('Workup',1)
add_p(doc,'Workup’a başvurma nedenim genel bir eğitim programına katılmak değil; İş Bankası ve iştirakleriyle gerçek bir uyum operasyonu problemini doğrulamak. Wardproof’un üçüncü taraf güvence, kritik hizmet, denetim kanıtı veya görevler ayrılığı modüllerinden birini 6–8 haftalık sınırlı PoC ile sınamak istiyorum. Program sonunda ölçülmüş bir vaka, doğru fiyat ve yatırım görüşmesine hazır bir ticari anlatı hedefliyorum.')
doc.add_heading('QNBEYOND',1)
add_p(doc,'QNB ile yapılabilecek en doğru başlangıcın küçük ve güvenli bir validasyon olduğunu düşünüyorum. Üretim sistemlerine yazmadan, sınırlı bir kontrol ve kanıt seti üzerinden mevcut hazırlık süresini ölçebilir; Wardproof ile oluşan karar ve denetim izini karşılaştırabiliriz. Başarılı olursa kapsamı üçüncü taraf riski veya kritik hizmet dayanıklılığına genişletebiliriz.')
doc.add_heading('Kısa ilk temas mesajı',1)
callout(doc,'E-POSTA / LINKEDIN','Merhaba, ben [AD SOYAD]. Finans kuruluşlarında mevzuat yükümlülüğünü kontrol, test ve kanıta bağlayan Wardproof.com adlı bir B2B uyum ürünü geliştiriyorum. Çalışan bir MVP’m var; şu anda geniş bir satış talebi yerine 6–8 haftalık, sınırlı bir banka PoC’si arıyorum. Uygun görürseniz ürünü 20 dakikalık bir demoda gösterip hangi kullanım alanının kurumunuz için anlamlı olabileceğini konuşmak isterim.')

page(doc); add_title(doc,'8. Sık sorulan zor sorular')
qas=[
('Müşteriniz var mı?','Henüz doğrulanmış ücretli müşterim yok. Ürünü teknik olarak çalışır hale getirdim; şimdi iki–üç tasarım ortağıyla dar pilotlar üzerinden ödeme isteğini ve kullanım değerini doğrulamak istiyorum.'),
('Neden mevcut GRC ürünleri yetmiyor?','Mevcut ürünlerin yetersiz olduğunu genellemek doğru olmaz. Wardproof’un iddiası daha dar: yerel yükümlülüğü kontrolün gerçek testine ve kanıtına bağlamak, sonuç durumlarını birbirine karıştırmamak ve kapanışı yeniden testle doğrulamak.'),
('Yapay zekâ ne yapıyor?','Yapay zekâ uyum durumuna veya puana karar vermiyor. Yardımcı özet, taslak ve sınıflandırma için kullanılabilir; kaynak, insan onayı ve deterministik iş kuralları kararın önünde kalır.'),
('Hukuki sorumluluğu nasıl yönetiyorsunuz?','Wardproof hukuki görüş veya uyum garantisi sunmaz. Kaynak sürümü, uzman doğrulaması, kurum kararı ve platform kontrolü ayrı gösterilir. Belirsiz içerik doğrulanmış gibi işaretlenmez.'),
('Tek kurucu riski?','Bu riskin farkındayım. Teknik ürün geliştirmeyi ben yürütüyorum; regülasyon doğrulaması ve kurumsal satış tarafında tamamlayıcı ekip kurmak istiyorum. Program desteğini bu nedenle yalnız finansman olarak görmüyorum.'),
('Gelir modeliniz doğrulandı mı?','Hayır. Yıllık B2B SaaS ve kapsamlı pilot modeli planlanıyor; fakat fiyat ve dönüşüm verisi müşteriyle doğrulanmadan kesinmiş gibi sunulmayacak.'),
('Veri güvenliği nasıl sağlanıyor?','Tenant izolasyonu veritabanı politikalarıyla uygulanıyor; kritik kayıtlar append-only veya immutable kurallarla korunuyor; kimlik atfı kullanıcı girdisine bırakılmıyor. Kurumsal üretim öncesinde bağımsız pentest ve restore provası tamamlanacak.'),
('Neden şimdi yatırım?','Yeni özellikleri çoğaltmak için değil; şirketleşme, güvenlik doğrulaması, uzman hukuk/regülasyon kontrolü ve ilk pilotları aynı anda yürütebilmek için. Teknik çekirdek var; ticari doğrulamanın önü açılmalı.')]
for q,a in qas:
    doc.add_heading(q,2); add_p(doc,a)

page(doc); add_title(doc,'9. Gönderim kontrol listesi')
for x in ['Başvuru formundaki tüm rakamların kaynağı var','Wardproof.com adı her belgede aynı yazılıyor','KALKAN_OS yalnız teknik kod adı olarak açıklanıyor','Demo bağlantısı başvuru dışından açılıyor','Demo verisi gerçek müşteri verisi içermiyor','Kurucu bilgileri ve LinkedIn güncel','Şirketleşme durumu dürüstçe belirtilmiş','BiGG öncesinde ortaklık/şirket uygunluğu yazılı teyit edilmiş','Müşteri/gelir yoksa varmış gibi yazılmamış','Tek pilot alanı ve tek başarı ölçütü seçilmiş','PDF/DOCX dosya adları sade ve Türkçe karakter sorunsuz','Son gönderimden önce programın güncel takvimi kontrol edilmiş']:
    bullet(doc,'[ ] '+x)
doc.add_heading('Önerilen ekler',1)
for x in ['Wardproof_Kisa_Sunum.pptx','Wardproof_Basvuru_Ana_Cevap_Bankasi.docx','İki dakikalık demo videosu bağlantısı','Kurucu özgeçmişi (1 sayfa)','Varsa LOI / pilot niyet mektubu','Güncel test ve canlı sistem doğrulama özeti']:
    bullet(doc,x)
callout(doc,'SON NOT','Bu metinler bilinçli olarak abartısız yazıldı. Başvuru değerlendirmesinde ürünün teknik derinliği zaten güçlü; güveni artıracak şey süslü ifade değil, neyin çalıştığını ve neyin henüz doğrulanmadığını açıkça ayırmak olacaktır.')

path=OUT/'Wardproof_Basvuru_Ana_Cevap_Bankasi.docx'
doc.save(path)

# One-pager
one=Document(); setup(one)
add_title(one,'WARDPROOF.COM','Finans kuruluşları için kanıta dayalı sürekli uyum')
callout(one,'KISA TANIM','Wardproof, mevzuat yükümlülüğünü kontrol, test, kanıt, bulgu ve bağımsız kapanış adımlarına bağlayan bir B2B regtech platformudur.')
one.add_heading('Sorun',1); add_p(one,'Uyum ekipleri farklı sistem, tablo ve klasörlerde çalışıyor. Kontrolün var olması ile gerçekten çalıştığının test edilmesi ve denetçiye savunulabilir kanıt sunulması arasında kopukluk oluşuyor.')
one.add_heading('Çözüm',1); add_p(one,'Wardproof her uyum iddiasının kaynağını, uygulanabilirlik kararını, kontrolünü, testini ve kanıtını tek zincirde tutuyor. “Başarısız”, “ölçülemedi”, “eski” ve “istisna kabul edildi” durumlarını ayırıyor; kritik bulguların başarılı yeniden test olmadan kapanmasını engelliyor.')
one.add_heading('Bugün nerede?',1)
for x in ['wardproof.com üzerinden erişilebilen çalışan ürün','Çok kiracılı veri izolasyonu ve denetim izi','Kontrol testi, bulgu/retest, SoD, üçüncü taraf ve Proof Room akışları','Yoğun birim/RLS ve gerçek tarayıcı doğrulaması','Henüz doğrulanmış ücretli müşteri ve fiyat yok']:
    bullet(one,x)
one.add_heading('İlk pilot',1); add_p(one,'Tek kurumda 20–30 kritik kontrol veya tek üçüncü taraf güvence paketi. Altı hafta içinde mevcut hazırlık süresi ölçülür, kanıt zinciri kurulur ve pilot sonunda doğrulanabilir yönetim özeti üretilir.')
one.add_heading('Aradığım destek',1); add_p(one,'Şirketleşme ve ilk pilot kaynağı; finans kuruluşuna erişim; regülasyon uzmanlığı; B2B satış mentorluğu; hukuk, muhasebe ve güvenlik doğrulaması.')
one.add_heading('İletişim',1); add_p(one,'[AD SOYAD]  •  [E-POSTA]  •  [TELEFON]  •  wardproof.com')
one.save(OUT/'Wardproof_One_Pager.docx')

# Founder profile template
cv=Document(); setup(cv); add_title(cv,'[AD SOYAD]','Wardproof.com Kurucusu')
add_p(cv,'[ŞEHİR] • [E-POSTA] • [TELEFON] • [LINKEDIN]')
cv.add_heading('Kısa profil',1); add_p(cv,'Finansal uyum, siber güvenlik ve yazılım ürünleri kesişiminde çalışan bir girişimciyim. Wardproof’u finans kuruluşlarının mevzuat, kontrol, test ve kanıt süreçlerini izlenebilir hale getirmek için geliştiriyorum. [MESLEKİ GEÇMİŞİNİZİ 2–3 CÜMLEYLE EKLEYİN.]')
cv.add_heading('Wardproof.com',1); add_p(cv,'Kurucu • [BAŞLANGIÇ AY/YIL] – devam')
for x in ['Çok kiracılı B2B uyum platformunun ürün ve teknik geliştirmesi','Kontrol, kanıt, bulgu, retest ve bağımsız onay iş akışları','Veritabanı seviyesinde yetkilendirme, audit ve değiştirilemezlik kuralları','Canlı dağıtım, otomatik test ve gerçek tarayıcı doğrulama süreçleri']:
    bullet(cv,x)
cv.add_heading('Deneyim',1); add_p(cv,'[ŞİRKET / ROL / TARİH]'); add_p(cv,'[En ilgili sorumluluk ve ölçülebilir sonuçlar.]')
cv.add_heading('Eğitim',1); add_p(cv,'[OKUL / BÖLÜM / YIL]')
cv.add_heading('Yetkinlikler',1); add_p(cv,'Ürün geliştirme • B2B SaaS • Next.js/TypeScript • PostgreSQL/Supabase • Güvenli çok kiracılı mimari • [DİĞER]')
cv.save(OUT/'Wardproof_Kurucu_Profil_Sablonu.docx')

print(str(OUT.resolve()))
